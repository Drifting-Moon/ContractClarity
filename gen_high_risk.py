from docx import Document

doc = Document()
doc.add_heading('High Risk Contract', 0)
doc.add_paragraph('This contract includes several high risk terms for testing.')
doc.add_paragraph('1. The Provider may terminate this agreement without cause at any time.')
doc.add_paragraph('2. The User agrees to indemnify the Provider against all claims.')
doc.add_paragraph('3. This agreement shall automatically renew for successive terms.')
doc.add_paragraph('4. The Provider shall have unlimited liability.')
doc.add_paragraph('5. Any disputes shall be settled by binding arbitration.')
doc.save('high_risk_contract.docx')
print("high_risk_contract.docx created")
