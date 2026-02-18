from docx import Document

doc = Document()
doc.add_heading('Contract Agreement', 0)
doc.add_paragraph('This is a sample contract for testing purposes.')
doc.add_paragraph('Clause 1: termination. This agreement can be terminated with 30 days notice.')
doc.save('test_contract.docx')
print("test_contract.docx created")
